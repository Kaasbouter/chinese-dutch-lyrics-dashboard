from __future__ import annotations

import io

import pytest

from lyrics_dashboard.alignment import build_exact_manual_plan
from lyrics_dashboard.converter import (
    ConversionSettings,
    convert_lyrics,
    encode_utf8_txt,
)
from lyrics_dashboard.extractors import extract_text
from lyrics_dashboard.models import AlignedLine, TranslationReference
from lyrics_dashboard.parser import parse_lyrics


def _single_language_source(leading_link: str) -> str:
    return (
        f"{leading_link}\n\n"
        "1 Grace\n\n"
        "Verse 1\n"
        "Lyrics remain unchanged.\n"
    )


@pytest.mark.parametrize(
    "leading_link",
    [
        "https://example.com/forbiddenmarker",
        "   http://example.com/forbiddenmarker   ",
        "www.example.com/forbiddenmarker",
        "ftp://example.com/forbiddenmarker",
        "example.com/forbiddenmarker",
        "youtube.com/watch/forbiddenmarker",
        "https://youtu.be/forbiddenmarker",
        "bit.ly/forbiddenmarker",
        "https://not-a-real-host.invalid/forbiddenmarker",
    ],
    ids=[
        "https",
        "http",
        "www",
        "ftp",
        "standalone-domain",
        "youtube",
        "youtube-short",
        "shortened-url",
        "inaccessible-host",
    ],
)
def test_supported_leading_url_forms_never_reach_single_language_output(
    leading_link: str,
) -> None:
    extracted = extract_text(
        "single-language.txt",
        _single_language_source(leading_link).encode("utf-8"),
    )
    parsed = parse_lyrics(extracted)
    preview = convert_lyrics(parsed, None, ConversionSettings())
    downloaded = encode_utf8_txt(preview)

    assert extracted.startswith("1 Grace\n")
    assert parsed.raw_title == "Grace"
    assert parsed.dutch_title == "Grace"
    assert parsed.sections[0].lines == ("Lyrics remain unchanged.",)
    assert "forbiddenmarker" not in repr(parsed).lower()
    assert "forbiddenmarker" not in preview.lower()
    assert downloaded.decode("utf-8") == preview
    assert b"forbiddenmarker" not in downloaded.lower()


def test_all_initial_blanks_multiple_links_and_wrapped_fragments_are_removed() -> None:
    source = (
        "\n\n"
        "https://example.com/first-forbidden\n"
        "\n\n\n"
        "www.youtube.com/watch/second-forbidden\n"
        "\n"
        "https://example.com/very-long-\n"
        "song-page-forbidden\n"
        "\n\n\n"
        "1 Song title\n\n"
        "Verse 1\n"
        "Lyrics stay in place\n"
    )

    extracted = extract_text("wrapped-links.txt", source.encode("utf-8"))
    parsed = parse_lyrics(extracted)

    assert extracted.splitlines()[0] == "1 Song title"
    assert parsed.raw_title == "Song title"
    assert parsed.sections[0].lines == ("Lyrics stay in place",)
    for fragment in (
        "example.com",
        "youtube.com",
        "first-forbidden",
        "second-forbidden",
        "song-page-forbidden",
    ):
        assert fragment not in extracted
        assert fragment not in repr(parsed)


def test_link_after_explicit_title_heading_is_removed_before_title_content() -> None:
    parsed = parse_lyrics(
        "[Title]\n"
        "https://example.com/forbidden-title-link\n\n"
        "Song title\n\n"
        "Verse 1\n"
        "Lyrics\n"
    )

    assert parsed.raw_title == "Song title"
    assert parsed.dutch_title == "Song title"
    assert parsed.sections[0].lines == ("Lyrics",)
    assert "forbidden-title-link" not in repr(parsed)


def test_bilingual_parsing_mapping_preview_and_txt_exclude_leading_link() -> None:
    source = (
        "https://example.com/forbidden-bilingual-marker\n\n"
        "中文歌名 | Nederlandse titel\n\n"
        "Verse 1\n"
        "中文歌词保留\n\n"
        "Verse 2\n"
        "Nederlandse regel blijft\n"
    )
    parsed = parse_lyrics(extract_text("bilingual.txt", source.encode("utf-8")))
    plan = build_exact_manual_plan(
        parsed,
        {0: [1]},
        {
            0: (
                AlignedLine(
                    source_line_indices=(0,),
                    translation_references=(TranslationReference(1, (0,)),),
                ),
            ),
        },
    )
    preview = convert_lyrics(parsed, plan, ConversionSettings(switch_index=1))
    downloaded = encode_utf8_txt(preview)

    assert parsed.chinese_title == "中文歌名"
    assert parsed.dutch_title == "Nederlandse titel"
    assert [section.lines for section in parsed.sections] == [
        ("中文歌词保留",),
        ("Nederlandse regel blijft",),
    ]
    assert [section.original_index for section in parsed.sections] == [0, 1]
    assert "forbidden-bilingual-marker" not in repr(parsed)
    assert "forbidden-bilingual-marker" not in repr(plan)
    assert "forbidden-bilingual-marker" not in preview
    assert downloaded.decode("utf-8") == preview
    assert b"forbidden-bilingual-marker" not in downloaded


def test_dotted_title_inline_domain_and_later_url_are_not_removed() -> None:
    source = (
        "Song.Title\n\n"
        "Verse 1\n"
        "This sentence has a period.\n"
        "Visit example.com today\n"
        "https://example.com/lyric-link\n"
    )

    extracted = extract_text("ordinary-periods.txt", source.encode("utf-8"))
    parsed = parse_lyrics(extracted)

    assert extracted == source
    assert parsed.raw_title == "Song.Title"
    assert parsed.sections[0].lines == (
        "This sentence has a period.",
        "Visit example.com today",
        "https://example.com/lyric-link",
    )


@pytest.mark.parametrize(
    ("leading_link", "title"),
    [
        ("https://example.com/", "Song.Title"),
        ("https://example.com/complete-", "Grace"),
    ],
)
def test_complete_url_does_not_consume_adjacent_punctuation_or_plain_title(
    leading_link: str,
    title: str,
) -> None:
    source = (
        f"{leading_link}\n"
        f"{title}\n\n"
        "Verse 1\n"
        "Lyrics remain\n"
    )

    extracted = extract_text("dotted-title.txt", source.encode("utf-8"))
    parsed = parse_lyrics(extracted)

    assert extracted.startswith(f"{title}\n")
    assert parsed.raw_title == title
    assert parsed.sections[0].lines == ("Lyrics remain",)


def _add_docx_hyperlink(paragraph, display_text: str, destination: str) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(
        destination,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = display_text
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_docx_field_hyperlink(paragraph, display_text: str, destination: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)

    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' HYPERLINK "{destination}" '
    instruction_run.append(instruction)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)

    display_run = OxmlElement("w:r")
    display = OxmlElement("w:t")
    display.text = display_text
    display_run.append(display)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)

    for element in (
        begin_run,
        instruction_run,
        separate_run,
        display_run,
        end_run,
    ):
        paragraph._p.append(element)


def test_docx_leading_hyperlink_display_and_destination_are_excluded() -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("")
    leading_hyperlink = document.add_paragraph()
    _add_docx_hyperlink(
        leading_hyperlink,
        "Open song website forbidden display",
        "https://hidden.invalid/forbidden-destination",
    )
    field_hyperlink = document.add_paragraph()
    _add_docx_field_hyperlink(
        field_hyperlink,
        "Click here forbidden field display",
        "https://hidden.invalid/forbidden-field-destination",
    )
    document.add_paragraph("")
    document.add_paragraph("1 Grace")
    document.add_paragraph("Verse 1")
    document.add_paragraph("Ordinary lyric line")
    later_hyperlink = document.add_paragraph()
    _add_docx_hyperlink(
        later_hyperlink,
        "Later hyperlink lyric stays",
        "https://hidden.invalid/later-destination",
    )
    buffer = io.BytesIO()
    document.save(buffer)

    extracted = extract_text("hyperlinks.docx", buffer.getvalue())
    parsed = parse_lyrics(extracted)
    preview = convert_lyrics(parsed, None, ConversionSettings())

    assert extracted.splitlines() == [
        "1 Grace",
        "Verse 1",
        "Ordinary lyric line",
        "Later hyperlink lyric stays",
    ]
    assert parsed.raw_title == "Grace"
    assert parsed.sections[0].lines == (
        "Ordinary lyric line",
        "Later hyperlink lyric stays",
    )
    assert "forbidden display" not in repr(parsed)
    assert "forbidden field display" not in repr(parsed)
    assert "forbidden-destination" not in extracted
    assert "forbidden-field-destination" not in extracted
    assert "forbidden-destination" not in preview
    assert "forbidden-field-destination" not in preview
    assert "Later hyperlink lyric stays" in preview
