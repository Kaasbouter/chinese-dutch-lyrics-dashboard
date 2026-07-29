from __future__ import annotations

import hashlib
import io
from pathlib import Path

import pytest

from lyrics_dashboard.alignment import (
    build_exact_manual_plan,
    parse_line_spec,
    parse_manual_rows,
)
from lyrics_dashboard.converter import ConversionSettings, convert_lyrics
from lyrics_dashboard.drag_mapping import (
    build_drag_line_groups,
    decode_drag_containers,
    make_chinese_line_groups,
)
from lyrics_dashboard.errors import PairingError
from lyrics_dashboard.extractors import extract_text
from lyrics_dashboard.models import AlignedLine, TranslationReference
from lyrics_dashboard.parser import parse_lyrics
from lyrics_dashboard.splitter import split_lyric

SOURCE = """1 祢信實何廣大 Groot is Uw trouw O Heer

Verse 1
祢信實何廣大哦神我的父
在祢沒有絲毫轉動影兒

Chorus 1
祢信實何廣大祢信實何廣大

Verse 2
祢不改變祢憐愛也無減少
像祢以往一樣直到永遠

Verse 4
Groot is Uw trouw, o Heer mijn God en Vader
Er is geen schaduw van omkeer bij U
Ben ik ontrouw, Gij blijft immer Dezelfde

Chorus 2
Groot is Uw trouw, o Heer, groot is Uw trouw, o Heer
"""


def _exact_plan(parsed):
    selections = {
        0: [3],
        1: [4],
        2: [3],
    }
    groups = {
        0: (
            AlignedLine((0,), (TranslationReference(3, (0,)),)),
            AlignedLine((1,), (TranslationReference(3, (1,)),)),
        ),
        1: (
            AlignedLine((0,), (TranslationReference(4, (0,)),)),
        ),
        2: (
            AlignedLine((0, 1), (TranslationReference(3, (2,)),)),
        ),
    }
    return build_exact_manual_plan(parsed, selections, groups)


def test_parser_accepts_unequal_language_section_counts() -> None:
    parsed = parse_lyrics(SOURCE)

    assert parsed.chinese_title == "祢信實何廣大"
    assert parsed.dutch_title == "Groot is Uw trouw O Heer"
    assert [section.language for section in parsed.sections] == ["zh", "zh", "zh", "nl", "nl"]
    assert any("3 Chinese sections and 2 Dutch sections" in warning for warning in parsed.warnings)


def test_parser_accepts_explicit_title_heading() -> None:
    parsed = parse_lyrics(
        "[Title]\n"
        "\u4e2d\u6587\u6b4c\u540d | Nederlandse titel\n\n"
        "Verse 1\n"
        "\u4e2d\u6587\u6b4c\u8bcd\n\n"
        "Verse 2\n"
        "Nederlandse liedtekst\n"
    )

    assert parsed.chinese_title == "\u4e2d\u6587\u6b4c\u540d"
    assert parsed.dutch_title == "Nederlandse titel"
    assert [section.label for section in parsed.sections] == ["Verse 1", "Verse 2"]


def test_exact_manual_alignment_supports_unequal_sections_and_lines() -> None:
    parsed = parse_lyrics(SOURCE)
    plan = _exact_plan(parsed)

    chinese_verse = plan.for_section(2)
    dutch_verse = plan.for_section(3)
    assert chinese_verse.counterpart_section_indices == (3,)
    assert chinese_verse.aligned_lines[0].source_line_indices == (0, 1)
    assert dutch_verse.counterpart_section_indices == (0, 2)
    assert dutch_verse.aligned_lines[-1].translation_references[0].line_indices == (0, 1)


def test_converter_switches_order_at_selected_original_section() -> None:
    parsed = parse_lyrics(SOURCE)
    output = convert_lyrics(
        parsed,
        _exact_plan(parsed),
        ConversionSettings(switch_index=3, chinese_max_length=10, dutch_max_length=40),
    )

    assert "[Title]\n祢信實何廣大 Groot is Uw trouw O Heer" in output
    first_verse = output.split("[Verse 1]\n", 1)[1].split("\n\n", 1)[0]
    dutch_verse = output.split("[Verse 4]\n", 1)[1].split("\n\n", 1)[0]
    assert first_verse.splitlines()[0].startswith("祢信實")
    assert "|Groot is Uw trouw" in first_verse.splitlines()[0]
    assert dutch_verse.splitlines()[0].startswith("Groot is Uw trouw")
    assert "|祢信實" in dutch_verse.splitlines()[0]


def test_representative_bilingual_output_remains_byte_for_byte_unchanged() -> None:
    parsed = parse_lyrics(SOURCE)
    output = convert_lyrics(
        parsed,
        _exact_plan(parsed),
        ConversionSettings(switch_index=3, chinese_max_length=10, dutch_max_length=40),
    )

    assert hashlib.sha256(output.encode("utf-8")).hexdigest() == (
        "fadb1dfc2dba342677582a927c809382cdd79f4c495d79bccc3fd917b11ce1c6"
    )


def test_manual_alignment_requires_each_dutch_section_to_be_covered() -> None:
    parsed = parse_lyrics(SOURCE)
    groups = {
        0: (AlignedLine((0, 1), (TranslationReference(3, (0, 1)),)),),
        1: (AlignedLine((0,), (TranslationReference(3, (2,)),)),),
        2: (AlignedLine((0, 1), (TranslationReference(3, (0, 1)),)),),
    }
    with pytest.raises(PairingError, match="Every Dutch section must be selected"):
        build_exact_manual_plan(parsed, {0: [3], 1: [3], 2: [3]}, groups)


def test_manual_alignment_requires_every_dutch_line_to_be_used() -> None:
    parsed = parse_lyrics(SOURCE)
    groups = {
        0: (
            AlignedLine((0,), (TranslationReference(3, (0,)),)),
            AlignedLine((1,), (TranslationReference(3, (1,)),)),
        ),
        1: (AlignedLine((0,), (TranslationReference(4, (0,)),)),),
        2: (AlignedLine((0, 1), (TranslationReference(3, (1,)),)),),
    }
    with pytest.raises(PairingError, match=r"Dutch \[Verse 4\] line 3"):
        build_exact_manual_plan(parsed, {0: [3], 1: [4], 2: [3]}, groups)


def test_parse_manual_rows_supports_ranges_and_multiple_sections() -> None:
    parsed = parse_lyrics(
        """Verse 1
中文第一行
中文第二行

Verse 2
Dutch first line
Dutch second line

Bridge
Dutch bridge line
"""
    )
    rows = [
        {"Chinese line(s)": "1", "Dutch reference(s)": "D1:1-2"},
        {"Chinese line(s)": "2", "Dutch reference(s)": "D2:1"},
    ]
    groups = parse_manual_rows(parsed, 0, [1, 2], rows, {"D1": 1, "D2": 2})

    assert groups[0].source_line_indices == (0,)
    assert groups[0].translation_references[0].line_indices == (0, 1)
    assert groups[1].translation_references[0].section_index == 2


def test_parse_line_spec_rejects_nonconsecutive_list_syntax() -> None:
    with pytest.raises(PairingError, match="line number or range"):
        parse_line_spec("1,3", 4, "Example")


def test_drag_grouping_joins_consecutive_chinese_lines() -> None:
    assert make_chinese_line_groups(4, (False, True, True)) == (
        (0,),
        (1, 2, 3),
    )


def test_drag_board_decodes_unique_cards_without_parsing_lyric_text() -> None:
    token_to_reference = {
        "D1:1 — Repeated lyric": (3, 0),
        "D1:2 — Repeated lyric": (3, 1),
    }
    unassigned, assignments = decode_drag_containers(
        [
            {"header": "Unassigned", "items": ["D1:2 — Repeated lyric"]},
            {"header": "Chinese 1", "items": ["D1:1 — Repeated lyric"]},
        ],
        unassigned_header="Unassigned",
        target_headers=("Chinese 1",),
        token_to_reference=token_to_reference,
    )

    assert unassigned == ((3, 1),)
    assert assignments == (((3, 0),),)


def test_drag_board_rejects_lost_cards() -> None:
    with pytest.raises(PairingError, match="lost or added Dutch cards"):
        decode_drag_containers(
            [
                {"header": "Unassigned", "items": []},
                {"header": "Chinese 1", "items": ["D1:1 — First"]},
            ],
            unassigned_header="Unassigned",
            target_headers=("Chinese 1",),
            token_to_reference={
                "D1:1 — First": (3, 0),
                "D1:2 — Second": (3, 1),
            },
        )


def test_drag_groups_support_many_to_one_with_shared_lines_unassigned() -> None:
    parsed = parse_lyrics(SOURCE)
    groups = build_drag_line_groups(
        parsed,
        source_section_index=0,
        selected_dutch_sections=(3,),
        source_line_groups=((0, 1),),
        assignments=(((3, 0),),),
        unassigned=((3, 1), (3, 2)),
    )

    assert groups[0].source_line_indices == (0, 1)
    assert groups[0].translation_references == (
        TranslationReference(section_index=3, line_indices=(0,)),
    )


def test_drag_groups_preserve_runs_across_multiple_dutch_sections() -> None:
    parsed = parse_lyrics(SOURCE)
    groups = build_drag_line_groups(
        parsed,
        source_section_index=0,
        selected_dutch_sections=(3, 4),
        source_line_groups=((0,), (1,)),
        assignments=(
            ((3, 0), (3, 1), (4, 0)),
            ((3, 2),),
        ),
        unassigned=(),
    )

    assert groups[0].translation_references == (
        TranslationReference(section_index=3, line_indices=(0, 1)),
        TranslationReference(section_index=4, line_indices=(0,)),
    )
    assert groups[1].translation_references == (
        TranslationReference(section_index=3, line_indices=(2,)),
    )


def test_drag_groups_reject_dutch_cards_out_of_order() -> None:
    parsed = parse_lyrics(SOURCE)
    with pytest.raises(PairingError, match="must stay in line order"):
        build_drag_line_groups(
            parsed,
            source_section_index=0,
            selected_dutch_sections=(3,),
            source_line_groups=((0,), (1,)),
            assignments=(
                ((3, 1),),
                ((3, 0),),
            ),
            unassigned=((3, 2),),
        )


def test_chinese_fallback_split_never_exceeds_limit() -> None:
    result = split_lyric("祢信實何廣大哦神我的父", "zh", 10)
    parts = result.split("//")
    assert len(parts) == 2
    assert all(len(part) <= 10 for part in parts)
    assert "".join(parts) == "祢信實何廣大哦神我的父"


def test_chinese_split_normalizes_internal_whitespace_without_losing_content() -> None:
    source = "\u4f60 \u597d\tAB"

    assert split_lyric(source, "zh", 40) == "\u4f60 \u597d AB"


def test_dutch_fallback_does_not_split_words() -> None:
    original = "Een bijzonder lange Nederlandse liedtekst zonder leestekens"
    result = split_lyric(original, "nl", 24)
    assert result.replace("//", " ") == original
    for segment in result.split("//"):
        assert not segment.startswith(" ")
        assert not segment.endswith(" ")


def test_text_extraction_distinguishes_cp1252_from_utf16() -> None:
    cp1252_source = "\u2019a"
    utf16_source = "Verse 1\nHello"

    assert extract_text("lyrics.txt", cp1252_source.encode("cp1252")) == cp1252_source
    assert extract_text("lyrics.txt", utf16_source.encode("utf-16-le")) == utf16_source


def test_docx_extraction_preserves_paragraph_and_table_order() -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("Before table")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Inside table"
    document.add_paragraph("After table")
    buffer = io.BytesIO()
    document.save(buffer)

    assert extract_text("lyrics.docx", buffer.getvalue()).splitlines() == [
        "Before table",
        "Inside table",
        "After table",
    ]


def test_uploaded_basic_format_example_converts_end_to_end() -> None:
    sample = Path(__file__).resolve().parents[1] / "samples" / "basic_format_example.docx"
    parsed = parse_lyrics(extract_text(sample.name, sample.read_bytes()))
    # These constants encode the manually verified translations in the supplied sample.
    confirmed_pairs = {
        0: (4, ((0, 0), (1, 1), (2, 2), (3, 3))),
        1: (5, ((0, 0), (1, 1), (2, 2), (3, 3))),
        2: (6, ((0, 0), (1, 1), (2, 2), (3, 3))),
        3: (7, ((0, 0), (1, 1), (2, 2), (3, 3))),
    }
    selections = {
        source_index: [target_index]
        for source_index, (target_index, _line_pairs) in confirmed_pairs.items()
    }
    line_groups = {
        source_index: tuple(
            AlignedLine(
                source_line_indices=(source_line,),
                translation_references=(
                    TranslationReference(target_index, (target_line,)),
                ),
            )
            for source_line, target_line in line_pairs
        )
        for source_index, (target_index, line_pairs) in confirmed_pairs.items()
    }
    plan = build_exact_manual_plan(parsed, selections, line_groups)
    output = convert_lyrics(parsed, plan, ConversionSettings(switch_index=4))

    assert len(parsed.sections) == 8
    assert output.startswith("[Title]\n祢信實何廣大 Groot is Uw trouw O Heer")
    assert "[Verse 1]\n祢信實何廣大" in output
    assert "[Verse 4]\nGroot is Uw trouw" in output
    assert output.endswith("\n")
